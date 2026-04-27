"""
Modern Resume DOCX Generator — uses python-docx (pure Python, no system deps).
Parses the AI's clean Markdown output and produces a styled Word document.
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_paragraph_border_bottom(paragraph, color="3366CC", size="6"):
    """Add a bottom border to a paragraph (used for section dividers)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _hex_rgb(hex_str):
    hex_str = hex_str.lstrip("#")
    return RGBColor(int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16))


def _parse_content(raw_text):
    """Same parser logic as in pdf_generator — returns structured blocks."""
    raw_text = re.sub(r'\\[a-zA-Z]+(\{[^}]*\})*', ' ', raw_text)
    raw_text = re.sub(r'[{}]', '', raw_text)

    blocks = []
    for line in raw_text.splitlines():
        line = line.strip()
        if not line:
            continue

        h_match = re.match(r'^#{1,3}\s+(.*)', line)
        if h_match:
            blocks.append({"type": "section", "text": h_match.group(1).upper()})
            continue

        if re.match(r'^[A-Z][A-Z\s&/]{2,30}$', line):
            blocks.append({"type": "section", "text": line})
            continue

        bold_match = re.match(r'^(?:[-•›]\s+|\*\s+)?\*\*(.+?)\*\*(.*)$', line)
        if bold_match:
            inner = bold_match.group(1).strip()
            rest  = bold_match.group(2).strip().lstrip('–-|').strip()
            parts = re.split(r'\s*[|–-]\s*', inner + (" | " + rest if rest else ""))
            parts = [p.strip() for p in parts if p.strip()]
            title = parts[0] if len(parts) > 0 else inner
            sub   = parts[1] if len(parts) > 1 else ""
            date  = parts[2] if len(parts) > 2 else ""
            blocks.append({"type": "subheading", "title": title, "sub": sub, "date": date})
            continue

        bullet_match = re.match(r'^[-•*›]\s+(.*)', line)
        if bullet_match:
            blocks.append({"type": "bullet", "text": bullet_match.group(1)})
            continue

        num_match = re.match(r'^\d+\.\s+(.*)', line)
        if num_match:
            blocks.append({"type": "bullet", "text": num_match.group(1)})
            continue

        blocks.append({"type": "body", "text": line})

    return blocks


def generate_resume_docx(data_dict, output_path="updated_resume"):
    """
    Generate a modern DOCX resume.
    data_dict keys: name, contact_info, content
    """
    doc = Document()

    # ---- Page margins ----
    for section in doc.sections:
        section.left_margin   = Inches(0.6)
        section.right_margin  = Inches(0.6)
        section.top_margin    = Inches(0.5)
        section.bottom_margin = Inches(0.5)

    NAVY   = _hex_rgb("1A1A4D")
    BLUE   = _hex_rgb("3366CC")
    GRAY   = _hex_rgb("555577")
    BLACK  = _hex_rgb("252525")

    # ---- Header ----
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name_para.add_run(data_dict.get("name", "Resume").upper())
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = NAVY

    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_raw = data_dict.get("contact_info", "").replace("|", "  ·  ")
    run = contact_para.add_run(contact_raw)
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY

    # Accent divider after header
    rule = doc.add_paragraph()
    _set_paragraph_border_bottom(rule, color="3366CC", size="12")
    rule.paragraph_format.space_after = Pt(4)

    # ---- Body blocks ----
    blocks = _parse_content(data_dict.get("content", ""))

    for block in blocks:
        btype = block["type"]

        if btype == "section":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(2)
            run = p.add_run(block["text"])
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = BLUE
            _set_paragraph_border_bottom(p, color="3366CC", size="6")

        elif btype == "subheading":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(1)
            # Title (bold) + date (right-aligned via tab) + sub (italic)
            run = p.add_run(block["title"])
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = BLACK
            if block.get("date"):
                p.add_run("\t")
                dr = p.add_run(block["date"])
                dr.font.size = Pt(9)
                dr.font.color.rgb = GRAY
            if block.get("sub"):
                sp = doc.add_paragraph()
                sp.paragraph_format.space_before = Pt(0)
                sp.paragraph_format.space_after  = Pt(2)
                sr = sp.add_run(block["sub"])
                sr.italic = True
                sr.font.size = Pt(9)
                sr.font.color.rgb = GRAY

        elif btype == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_after  = Pt(1)
            run = p.add_run(block["text"])
            run.font.size = Pt(9.5)
            run.font.color.rgb = BLACK

        else:  # body
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(block["text"])
            run.font.size = Pt(9.5)
            run.font.color.rgb = BLACK

    out_file = f"{output_path}.docx"
    doc.save(out_file)
    return out_file
